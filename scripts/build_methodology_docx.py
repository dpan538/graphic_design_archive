from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "Methodology_v0.md"
DOCX_PATH = ROOT / "Methodology_v0.docx"


def set_cell_text(paragraph, text):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 31, 31)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(4)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Methodology v0 · Modern Graphic Design History Archive Index"
    paragraph.style = doc.styles["Footer"]
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def add_horizontal_rule(paragraph):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_runs_with_inline_markup(paragraph, text):
    # Convert common inline Markdown into readable Word runs.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label}: {url}")
            run.font.color.rgb = RGBColor(5, 99, 193)
            run.underline = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    first_content = True
    in_front_matter = True

    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            add_runs_with_inline_markup(p, line[2:])
            add_horizontal_rule(p)
            continue

        if line.startswith("**Project:**") or line.startswith("**Working definition:**") or line.startswith("**Date:**") or line.startswith("**Status:**"):
            p = doc.add_paragraph()
            add_runs_with_inline_markup(p, line)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(85, 85, 85)
            continue

        if line.startswith("## "):
            if first_content:
                first_content = False
            p = doc.add_paragraph(style="Heading 1")
            add_runs_with_inline_markup(p, line[3:])
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_runs_with_inline_markup(p, line[4:])
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_inline_markup(p, line[2:])
            continue

        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_runs_with_inline_markup(p, re.sub(r"^\d+\. ", "", line))
            continue

        p = doc.add_paragraph()
        add_runs_with_inline_markup(p, line)

    add_footer(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
